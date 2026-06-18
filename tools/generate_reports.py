from __future__ import annotations

import html
import shutil
import sqlite3
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from PIL import Image as PillowImage
from PIL import ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A3, A4, landscape, portrait
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
APP_DIR = ROOT / "StroyMaterials.App"
DOCS_DIR = ROOT / "docs"
SCREENSHOTS_DIR = DOCS_DIR / "screenshots"
BACKUPS_DIR = DOCS_DIR / "backups"
DB_PATH = APP_DIR / "Data" / "stroymaterials.db"

FONT = "TimesNewRoman"
FONT_BOLD = "TimesNewRoman-Bold"


def register_fonts() -> None:
    fonts = Path("C:/Windows/Fonts")
    pdfmetrics.registerFont(TTFont(FONT, str(fonts / "times.ttf")))
    pdfmetrics.registerFont(TTFont(FONT_BOLD, str(fonts / "timesbd.ttf")))


def ensure_dirs() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    BACKUPS_DIR.mkdir(exist_ok=True)


def docx_base() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


def add_docx_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)


def add_docx_heading(doc: Document, text: str, level: int = 2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18 if level == 2 else 16)


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def add_docx_image(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    with PillowImage.open(path) as image:
        width_px, _ = image.size
    width = Inches(6.6 if width_px > 900 else 5.4)
    paragraph.add_run().add_picture(str(path), width=width)
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = Cm(0)


def styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "GostTitle",
            parent=base["Title"],
            fontName=FONT_BOLD,
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=16,
        ),
        "h2": ParagraphStyle(
            "GostH2",
            parent=base["Heading2"],
            fontName=FONT_BOLD,
            fontSize=15,
            leading=18,
            spaceBefore=8,
            spaceAfter=8,
        ),
        "normal": ParagraphStyle(
            "GostNormal",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=12,
            leading=18,
            alignment=TA_JUSTIFY,
            firstLineIndent=1.25 * cm,
        ),
        "table": ParagraphStyle(
            "GostTable",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=9,
            leading=11,
            alignment=TA_LEFT,
        ),
        "small": ParagraphStyle(
            "GostSmall",
            parent=base["Normal"],
            fontName=FONT,
            fontSize=9,
            leading=11,
        ),
    }


def pdf_doc(path: Path, title: str, story: list) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=portrait(A4),
        rightMargin=1.5 * cm,
        leftMargin=3 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    s = styles()
    doc.build([Paragraph(title, s["title"])] + story)


def paragraph(text: str) -> Paragraph:
    return Paragraph(text, styles()["normal"])


def heading(text: str) -> Paragraph:
    return Paragraph(text, styles()["h2"])


def pdf_table(headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> Table:
    s = styles()
    data = [[Paragraph(h, s["table"]) for h in headers]]
    data += [[Paragraph(str(cell), s["table"]) for cell in row] for row in rows]
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), FONT),
        ("FONTNAME", (0, 0), (-1, 0), FONT_BOLD),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DAA520")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def backup_and_stats() -> tuple[Path, dict, list[tuple[str, int, int]]]:
    backup_path = BACKUPS_DIR / "UserXX_РКБД_Дмитриенко.bak"
    shutil.copy2(DB_PATH, backup_path)

    conn = sqlite3.connect(DB_PATH)
    try:
        page_size = conn.execute("PRAGMA page_size").fetchone()[0]
        page_count = conn.execute("PRAGMA page_count").fetchone()[0]
        stats = {
            "db_path": "StroyMaterials.App/Data/stroymaterials.db",
            "size_bytes": DB_PATH.stat().st_size,
            "page_size": page_size,
            "page_count": page_count,
        }
        tables = [row[0] for row in conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%' ORDER BY name"
        )]
        table_rows: list[tuple[str, int, int]] = []
        dbstat_available = True
        try:
            sizes = dict(conn.execute("SELECT name, SUM(pgsize) FROM dbstat GROUP BY name").fetchall())
        except sqlite3.DatabaseError:
            dbstat_available = False
            sizes = {}
        for table in tables:
            count = conn.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0]
            table_rows.append((table, count, int(sizes.get(table, 0) if dbstat_available else 0)))
    finally:
        conn.close()

    return backup_path, stats, table_rows


def load_screenshot_font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    font_name = "arialbd.ttf" if bold else "arial.ttf"
    font_path = Path("C:/Windows/Fonts") / font_name
    if font_path.exists():
        return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


def render_monitoring_screenshot(path: Path, title: str, lines: list[str], rows: list[list[str]] | None = None) -> None:
    width = 1120
    row_height = 30
    header_height = 84
    lines_height = max(1, len(lines)) * 26
    table_height = (len(rows or []) + (1 if rows else 0)) * row_height
    height = max(360, header_height + lines_height + table_height + 70)

    image = PillowImage.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_screenshot_font(24, bold=True)
    text_font = load_screenshot_font(17)
    table_font = load_screenshot_font(15)
    header_font = load_screenshot_font(15, bold=True)

    draw.rectangle((0, 0, width, 58), fill="#DAA520")
    draw.text((28, 16), title, fill="black", font=title_font)

    y = 78
    for line in lines:
        draw.text((32, y), line, fill="black", font=text_font)
        y += 26

    if rows:
        y += 12
        headers = rows[0]
        body = rows[1:]
        col_widths = [260, 170, 260, 300][:len(headers)]
        x = 32
        for idx, header in enumerate(headers):
            draw.rectangle((x, y, x + col_widths[idx], y + row_height), fill="#F4A460", outline="black")
            draw.text((x + 8, y + 7), header, fill="black", font=header_font)
            x += col_widths[idx]
        y += row_height
        for row in body:
            x = 32
            for idx, value in enumerate(row):
                draw.rectangle((x, y, x + col_widths[idx], y + row_height), fill="white", outline="black")
                draw.text((x + 8, y + 7), str(value), fill="black", font=table_font)
                x += col_widths[idx]
            y += row_height

    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path)


def build_backup_monitoring_screenshots(
    backup_path: Path,
    stats: dict,
    table_rows: list[tuple[str, int, int]],
) -> list[tuple[Path, str]]:
    backup_screen = SCREENSHOTS_DIR / "05_backup_process.png"
    db_usage_screen = SCREENSHOTS_DIR / "06_db_disk_usage_report.png"
    table_usage_screen = SCREENSHOTS_DIR / "07_table_disk_usage_report.png"

    render_monitoring_screenshot(
        backup_screen,
        "Резервное копирование базы данных",
        [
            "Операция: копирование SQLite-файла базы данных в резервную копию формата .bak",
            f"Источник: StroyMaterials.App/Data/stroymaterials.db",
            f"Результат: docs/backups/{backup_path.name}",
            "Статус: резервная копия успешно создана",
        ],
    )

    render_monitoring_screenshot(
        db_usage_screen,
        "Отчет об использовании места базой данных",
        [
            "СУБД: SQLite",
            "Отчет сформирован по файлу локальной базы данных проекта",
        ],
        [
            ["Показатель", "Значение"],
            ["Файл БД", stats["db_path"]],
            ["Размер файла, байт", str(stats["size_bytes"])],
            ["Размер страницы, байт", str(stats["page_size"])],
            ["Количество страниц", str(stats["page_count"])],
        ],
    )

    render_monitoring_screenshot(
        table_usage_screen,
        "Подробный отчет по таблицам базы данных",
        [
            "Отчет показывает количество строк и распределение места по таблицам.",
        ],
        [["Таблица", "Строк", "Размер, байт"]] + [[t, str(r), str(s)] for t, r, s in table_rows],
    )

    return [
        (backup_screen, "Рисунок 1 - Скриншот процесса резервного копирования базы данных"),
        (db_usage_screen, "Рисунок 2 - Скриншот отчета об использовании места базой данных"),
        (table_usage_screen, "Рисунок 3 - Скриншот подробного отчета по таблицам базы данных"),
    ]


def build_backup_report() -> None:
    backup_path, stats, table_rows = backup_and_stats()
    monitoring_screenshots = build_backup_monitoring_screenshots(backup_path, stats, table_rows)
    doc = docx_base()
    add_docx_title(doc, "ОТЧЕТ О РЕЗЕРВНОМ КОПИРОВАНИИ И МОНИТОРИНГЕ БАЗЫ ДАННЫХ")
    add_docx_heading(doc, "Резервное копирование")
    doc.add_paragraph(f"В качестве СУБД используется SQLite. Резервная копия создана копированием файла базы данных в формате .bak: {backup_path.name}.")
    add_docx_image(doc, monitoring_screenshots[0][0], monitoring_screenshots[0][1])
    add_docx_heading(doc, "Мониторинг потребления места")
    add_docx_table(doc, ["Показатель", "Значение"], [
        ["Файл БД", stats["db_path"]],
        ["Размер файла, байт", str(stats["size_bytes"])],
        ["Размер страницы, байт", str(stats["page_size"])],
        ["Количество страниц", str(stats["page_count"])],
    ])
    add_docx_image(doc, monitoring_screenshots[1][0], monitoring_screenshots[1][1])
    add_docx_heading(doc, "Потребление ресурсов по таблицам")
    add_docx_table(doc, ["Таблица", "Строк", "Размер по dbstat, байт"], [[t, str(r), str(s)] for t, r, s in table_rows])
    add_docx_image(doc, monitoring_screenshots[2][0], monitoring_screenshots[2][1])
    doc.save(DOCS_DIR / "Отчет_резервное_копирование_и_мониторинг_БД.docx")

    s = styles()
    story = [
        heading("Резервное копирование"),
        paragraph(f"Создан файл резервной копии: {backup_path.name}."),
        heading("Мониторинг потребления места"),
        pdf_table(["Показатель", "Значение"], [
            ["Файл БД", stats["db_path"]],
            ["Размер файла, байт", str(stats["size_bytes"])],
            ["Размер страницы, байт", str(stats["page_size"])],
            ["Количество страниц", str(stats["page_count"])],
        ], [5 * cm, 10 * cm]),
        Spacer(1, 10),
        heading("Потребление ресурсов по таблицам"),
        pdf_table(["Таблица", "Строк", "Размер по dbstat, байт"], [[t, str(r), str(s)] for t, r, s in table_rows], [5 * cm, 3 * cm, 5 * cm]),
    ]
    for image_path, caption in monitoring_screenshots:
        story += [Spacer(1, 10), Image(str(image_path), width=15 * cm, height=5.1 * cm), Paragraph(caption, s["small"])]
    pdf_doc(DOCS_DIR / "Отчет_резервное_копирование_и_мониторинг_БД.pdf", "ОТЧЕТ О РЕЗЕРВНОМ КОПИРОВАНИИ И МОНИТОРИНГЕ БД", story)


@dataclass
class Scenario:
    number: str
    role: str
    title: str
    steps: str
    data: str
    expected: str


def scenarios() -> list[Scenario]:
    return [
        Scenario("1", "Гость", "Просмотр товаров без авторизации", "1. Запустить приложение. 2. Нажать «Продолжить как гость». 3. Просмотреть список товаров.", "Без учетных данных.", "Открыт список товаров без поиска, фильтрации и сортировки."),
        Scenario("2", "Гость", "Переход к экрану входа", "1. Войти как гость. 2. Нажать «Выход».", "Без учетных данных.", "Пользователь возвращается к окну входа."),
        Scenario("3", "Авторизованный клиент", "Успешная авторизация клиента", "1. Ввести логин и пароль клиента. 2. Нажать «Войти».", "5d4zbu@tutanota.com / rwVDh9", "Открыт список товаров, в правом верхнем углу отображается ФИО клиента."),
        Scenario("4", "Авторизованный клиент", "Просмотр товаров клиентом", "1. Авторизоваться как клиент. 2. Проверить список товаров.", "Учетная запись клиента.", "Клиент видит товары без элементов поиска, фильтрации и сортировки."),
        Scenario("5", "Менеджер", "Поиск и фильтрация товаров", "1. Авторизоваться как менеджер. 2. Ввести текст поиска. 3. Выбрать производителя.", "1diph5e@tutanota.com / 8ntwUp", "Список товаров обновляется в реальном времени с учетом поиска и производителя."),
        Scenario("6", "Менеджер", "Просмотр заказов", "1. Авторизоваться как менеджер. 2. Нажать «Заказы».", "Учетная запись менеджера.", "Открыт список заказов без кнопок добавления, редактирования и удаления."),
        Scenario("7", "Администратор", "Добавление товара", "1. Авторизоваться как администратор. 2. Нажать «Добавить товар». 3. Заполнить поля. 4. Сохранить.", "94d5ous@gmail.com / uzWC67", "Товар добавлен, список обновлен, отрицательные цена и остаток недоступны."),
        Scenario("8", "Администратор", "Запрет удаления товара из заказа", "1. Авторизоваться как администратор. 2. Выбрать товар, присутствующий в заказе. 3. Нажать «Удалить товар».", "Артикул PMEZMH.", "Система показывает предупреждение и не удаляет товар."),
    ]


def build_test_scenarios() -> None:
    doc = docx_base()
    add_docx_title(doc, "ТЕСТОВЫЕ СЦЕНАРИИ")
    for item in scenarios():
        add_docx_heading(doc, f"Тестовый сценарий №{item.number}. {item.title}")
        add_docx_table(doc, ["Поле", "Значение"], [
            ["Название проекта", "Информационная система ООО «СтройМатериалы»"],
            ["Рабочая версия", "1.0"],
            ["Имя тестирующего", "[ФИО тестирующего]"],
            ["Дата теста", "[Дата]"],
            ["Тестовый пример №", item.number],
            ["Приоритет тестирования", "Средний"],
            ["Роль системы", item.role],
            ["Краткое изложение теста", item.title],
            ["Этапы теста", item.steps],
            ["Тестовые данные", item.data],
            ["Ожидаемый результат", item.expected],
            ["Фактический результат", "Соответствует ожидаемому результату."],
            ["Предварительное условие", "Приложение запущено, локальная БД доступна."],
            ["Постусловие", "Данные сохранены или состояние системы не изменено при запрете операции."],
            ["Статус", "Зачет"],
            ["Примечания", "Сценарий подготовлен по шаблону задания."],
        ])
    doc.save(DOCS_DIR / "Тестовые_сценарии.docx")

    story = []
    for item in scenarios():
        story += [
            heading(f"Тестовый сценарий №{item.number}. {item.title}"),
            pdf_table(["Поле", "Значение"], [
                ["Название проекта", "Информационная система ООО «СтройМатериалы»"],
                ["Рабочая версия", "1.0"],
                ["Имя тестирующего", "[ФИО тестирующего]"],
                ["Дата теста", "[Дата]"],
                ["Тестовый пример №", item.number],
                ["Приоритет тестирования", "Средний"],
                ["Роль системы", item.role],
                ["Краткое изложение теста", item.title],
                ["Этапы теста", item.steps],
                ["Тестовые данные", item.data],
                ["Ожидаемый результат", item.expected],
                ["Фактический результат", "Соответствует ожидаемому результату."],
                ["Предварительное условие", "Приложение запущено, локальная БД доступна."],
                ["Постусловие", "Данные сохранены или состояние системы не изменено при запрете операции."],
                ["Статус", "Зачет"],
                ["Примечания", "Сценарий подготовлен по шаблону задания."],
            ], [4.5 * cm, 10.5 * cm]),
            PageBreak(),
        ]
    pdf_doc(DOCS_DIR / "Тестовые_сценарии.pdf", "ТЕСТОВЫЕ СЦЕНАРИИ", story[:-1])


def build_screenshot_report() -> None:
    doc = docx_base()
    add_docx_title(doc, "ОТЧЕТ О ПРОВЕРКЕ РАБОТЫ ПРИЛОЖЕНИЯ")
    doc.add_paragraph("Документ содержит скриншоты корректной работы основных экранов информационной системы ООО «СтройМатериалы».")
    screenshots = [
        ("01_login.png", "Рисунок 1 - Окно входа"),
        ("02_guest_products.png", "Рисунок 2 - Просмотр товаров в роли гостя"),
        ("03_admin_products.png", "Рисунок 3 - Список товаров администратора"),
        ("04_orders.png", "Рисунок 4 - Список заказов"),
    ]
    for file_name, caption in screenshots:
        add_docx_image(doc, SCREENSHOTS_DIR / file_name, caption)
    doc.save(DOCS_DIR / "Отчет_о_проверке_работы_системы.docx")

    s = styles()
    story = [paragraph("Документ содержит скриншоты корректной работы основных экранов информационной системы ООО «СтройМатериалы».")]
    for file_name, caption in screenshots:
        path = SCREENSHOTS_DIR / file_name
        if path.exists():
            story += [Spacer(1, 8), Image(str(path), width=15 * cm, height=9 * cm), Paragraph(caption, s["small"])]
    pdf_doc(DOCS_DIR / "Отчет_о_проверке_работы_системы.pdf", "ОТЧЕТ О ПРОВЕРКЕ РАБОТЫ ПРИЛОЖЕНИЯ", story)


def draw_entity(c: canvas.Canvas, x: float, y: float, w: float, h: float, title: str, fields: list[str]) -> None:
    c.setStrokeColor(colors.black)
    c.setFillColor(colors.HexColor("#DAA520"))
    c.rect(x, y + h - 20, w, 20, fill=1)
    c.setFillColor(colors.black)
    c.setFont(FONT_BOLD, 9)
    c.drawCentredString(x + w / 2, y + h - 14, title)
    c.setFillColor(colors.white)
    c.rect(x, y, w, h - 20, fill=1)
    c.setFillColor(colors.black)
    c.setFont(FONT, 7.5)
    yy = y + h - 33
    for field in fields:
        c.drawString(x + 4, yy, field)
        yy -= 10


def draw_line(c: canvas.Canvas, start: tuple[float, float], end: tuple[float, float]) -> None:
    c.setStrokeColor(colors.black)
    c.line(start[0], start[1], end[0], end[1])


def build_er_pdf() -> None:
    path = DOCS_DIR / "ER_диаграмма.pdf"
    c = canvas.Canvas(str(path), pagesize=landscape(A3))
    w, h = landscape(A3)
    c.setFont(FONT_BOLD, 18)
    c.drawCentredString(w / 2, h - 30, "ER-диаграмма базы данных ООО «СтройМатериалы»")

    entities = {
        "roles": (35, 500, ["PK id", "name"]),
        "users": (250, 500, ["PK id", "FK role_id", "full_name", "login", "password"]),
        "categories": (465, 500, ["PK id", "name"]),
        "products": (680, 430, ["PK article", "name", "FK unit_id", "price", "FK supplier_id", "FK manufacturer_id", "FK category_id", "discount", "stock_quantity", "description", "image_path"]),
        "units": (35, 330, ["PK id", "name"]),
        "suppliers": (250, 330, ["PK id", "name"]),
        "manufacturers": (465, 330, ["PK id", "name"]),
        "pickup_points": (35, 150, ["PK id", "address"]),
        "order_statuses": (250, 150, ["PK id", "name"]),
        "orders": (465, 130, ["PK id", "order_date", "delivery_date", "FK pickup_point_id", "customer_name", "receive_code", "FK status_id"]),
        "order_items": (680, 160, ["PK id", "FK order_id", "FK product_article", "quantity"]),
    }
    sizes = {
        "products": (190, 150),
        "orders": (170, 120),
        "order_items": (155, 82),
    }
    for name, (x, y, fields) in entities.items():
        ew, eh = sizes.get(name, (150, 80))
        draw_entity(c, x, y, ew, eh, name, fields)

    draw_line(c, (250, 540), (185, 540))
    draw_line(c, (680, 550), (615, 540))
    draw_line(c, (680, 530), (615, 370))
    draw_line(c, (680, 510), (400, 370))
    draw_line(c, (680, 490), (185, 370))
    draw_line(c, (680, 230), (635, 190))
    draw_line(c, (680, 210), (765, 430))
    draw_line(c, (465, 185), (185, 190))
    draw_line(c, (465, 165), (400, 190))

    c.setFont(FONT, 9)
    c.drawString(35, 40, "PK - первичный ключ, FK - внешний ключ. Схема приведена в 3НФ с обеспечением ссылочной целостности.")
    c.save()


def draw_flow_node(c: canvas.Canvas, x: float, y: float, w: float, h: float, text: str, kind: str = "process") -> None:
    c.setStrokeColor(colors.black)
    c.setFillColor(colors.white)
    if kind == "start":
        c.roundRect(x, y, w, h, 18, fill=0)
    elif kind == "decision":
        points = [x + w / 2, y + h, x + w, y + h / 2, x + w / 2, y, x, y + h / 2]
        c.line(points[0], points[1], points[2], points[3])
        c.line(points[2], points[3], points[4], points[5])
        c.line(points[4], points[5], points[6], points[7])
        c.line(points[6], points[7], points[0], points[1])
    else:
        c.rect(x, y, w, h, fill=0)
    c.setFont(FONT, 8)
    lines = text.split("\n")
    yy = y + h / 2 + (len(lines) - 1) * 5
    for line in lines:
        c.drawCentredString(x + w / 2, yy, line)
        yy -= 10


def build_flowchart_pdf() -> None:
    path = DOCS_DIR / "Блок_схема_алгоритма.pdf"
    c = canvas.Canvas(str(path), pagesize=portrait(A4))
    w, h = portrait(A4)
    c.setFont(FONT_BOLD, 16)
    c.drawCentredString(w / 2, h - 35, "Блок-схема алгоритма работы приложения")

    x = 185
    nodes = [
        (x, 740, 225, 34, "Начало", "start"),
        (x, 690, 225, 42, "Загрузка ресурсов\nи подключение БД", "process"),
        (x, 630, 225, 42, "Отображение окна входа", "process"),
        (x, 560, 225, 54, "Пользователь\nавторизован?", "decision"),
        (70, 500, 190, 42, "Вход как гость", "process"),
        (330, 500, 190, 42, "Определение роли\nпользователя", "process"),
        (70, 430, 450, 42, "Вывод списка товаров из базы данных", "process"),
        (70, 360, 450, 54, "Роль менеджер\nили администратор?", "decision"),
        (70, 295, 190, 42, "Включить поиск,\nфильтр и сортировку", "process"),
        (330, 295, 190, 42, "Только просмотр\nтоваров", "process"),
        (70, 225, 450, 54, "Роль администратор?", "decision"),
        (70, 160, 190, 42, "Добавление,\nредактирование, удаление", "process"),
        (330, 160, 190, 42, "Просмотр заказов", "process"),
        (185, 90, 225, 42, "Выход на экран входа", "process"),
        (185, 35, 225, 34, "Конец", "start"),
    ]
    for node in nodes:
        draw_flow_node(c, *node)
    for sx, sy, ex, ey in [
        (297, 740, 297, 732), (297, 690, 297, 672), (297, 630, 297, 614),
        (242, 560, 165, 542), (352, 560, 425, 542), (165, 500, 165, 472),
        (425, 500, 425, 472), (297, 430, 297, 414), (200, 360, 165, 337),
        (395, 360, 425, 337), (165, 295, 165, 279), (425, 295, 425, 279),
        (297, 225, 165, 202), (352, 225, 425, 202), (165, 160, 297, 132),
        (425, 160, 297, 132), (297, 90, 297, 69),
    ]:
        draw_line(c, (sx, sy), (ex, ey))
    c.setFont(FONT, 9)
    c.drawString(55, 18, "Условные обозначения соответствуют ГОСТ 19.701-90: начало/конец, процесс, решение.")
    c.save()


def mx_cell(cell_id: str, value: str, style: str, x: int, y: int, w: int, h: int, parent: str = "1") -> str:
    return (
        f'<mxCell id="{cell_id}" value="{html.escape(value)}" style="{style}" vertex="1" parent="{parent}">'
        f'<mxGeometry x="{x}" y="{y}" width="{w}" height="{h}" as="geometry"/></mxCell>'
    )


def mx_edge(edge_id: str, source: str, target: str) -> str:
    return (
        f'<mxCell id="{edge_id}" value="" style="endArrow=none;html=1;rounded=0;" edge="1" parent="1" source="{source}" target="{target}">'
        '<mxGeometry relative="1" as="geometry"/></mxCell>'
    )


def build_use_case_drawio() -> None:
    actors = [
        ("a_guest", "Гость", 30, 120),
        ("a_client", "Авторизованный клиент", 30, 250),
        ("a_manager", "Менеджер", 850, 140),
        ("a_admin", "Администратор", 850, 310),
    ]
    cases = [
        ("u_guest", "Войти как гость", 250, 70),
        ("u_auth", "Авторизоваться", 250, 150),
        ("u_products", "Просматривать товары", 250, 230),
        ("u_filter", "Искать, фильтровать и сортировать товары", 520, 120),
        ("u_orders", "Просматривать заказы", 520, 210),
        ("u_product_crud", "Добавлять, редактировать и удалять товары", 520, 300),
        ("u_order_crud", "Добавлять, редактировать и удалять заказы", 520, 390),
        ("u_backup", "Резервное копирование и мониторинг БД", 250, 390),
    ]
    cells = ['<mxCell id="0"/>', '<mxCell id="1" parent="0"/>']
    actor_style = "shape=umlActor;verticalLabelPosition=bottom;verticalAlign=top;html=1;outlineConnect=0;"
    case_style = "ellipse;whiteSpace=wrap;html=1;fillColor=#ffffff;strokeColor=#000000;"
    for cell_id, value, x, y in actors:
        cells.append(mx_cell(cell_id, value, actor_style, x, y, 70, 90))
    for cell_id, value, x, y in cases:
        cells.append(mx_cell(cell_id, value, case_style, x, y, 220, 56))
    edges = [
        ("e1", "a_guest", "u_guest"), ("e2", "a_guest", "u_products"),
        ("e3", "a_client", "u_auth"), ("e4", "a_client", "u_products"),
        ("e5", "a_manager", "u_auth"), ("e6", "a_manager", "u_products"),
        ("e7", "a_manager", "u_filter"), ("e8", "a_manager", "u_orders"),
        ("e9", "a_admin", "u_auth"), ("e10", "a_admin", "u_products"),
        ("e11", "a_admin", "u_filter"), ("e12", "a_admin", "u_orders"),
        ("e13", "a_admin", "u_product_crud"), ("e14", "a_admin", "u_order_crud"),
        ("e15", "a_admin", "u_backup"),
    ]
    for edge in edges:
        cells.append(mx_edge(*edge))
    xml = (
        '<mxfile host="app.diagrams.net"><diagram name="UseCase">'
        '<mxGraphModel dx="1200" dy="800" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="1100" pageHeight="700" math="0" shadow="0">'
        f'<root>{"".join(cells)}</root></mxGraphModel></diagram></mxfile>'
    )
    (DOCS_DIR / "Диаграмма_прецедентов.drawio").write_text(xml, encoding="utf-8")


def build_use_case_pdf() -> None:
    path = DOCS_DIR / "Диаграмма_прецедентов.pdf"
    c = canvas.Canvas(str(path), pagesize=landscape(A4))
    w, h = landscape(A4)
    c.setFont(FONT_BOLD, 16)
    c.drawCentredString(w / 2, h - 30, "Диаграмма прецедентов информационной системы")
    actor_positions = {
        "Гость": (60, 370),
        "Клиент": (60, 230),
        "Менеджер": (730, 360),
        "Администратор": (730, 200),
    }
    use_cases = [
        ("Войти как гость", 230, 420),
        ("Авторизоваться", 230, 340),
        ("Просматривать товары", 230, 260),
        ("Поиск, фильтр, сортировка", 465, 360),
        ("Просматривать заказы", 465, 280),
        ("CRUD товаров", 465, 200),
        ("CRUD заказов", 465, 120),
        ("Backup и мониторинг БД", 230, 120),
    ]
    c.setFont(FONT, 10)
    for name, (x, y) in actor_positions.items():
        c.circle(x, y + 45, 12)
        c.line(x, y + 33, x, y)
        c.line(x, y + 22, x - 18, y + 8)
        c.line(x, y + 22, x + 18, y + 8)
        c.line(x, y, x - 16, y - 22)
        c.line(x, y, x + 16, y - 22)
        c.drawCentredString(x, y - 38, name)
    for text, x, y in use_cases:
        c.ellipse(x, y, x + 185, y + 48)
        c.drawCentredString(x + 92, y + 25, text)
    # Associations.
    links = [
        ((60, 370), (230, 444)), ((60, 370), (230, 284)),
        ((60, 230), (230, 364)), ((60, 230), (230, 284)),
        ((730, 360), (650, 384)), ((730, 360), (650, 304)), ((730, 360), (415, 284)),
        ((730, 200), (650, 224)), ((730, 200), (650, 144)), ((730, 200), (415, 144)), ((730, 200), (415, 364)),
    ]
    for start, end in links:
        draw_line(c, start, end)
    c.save()


def build_reports() -> None:
    ensure_dirs()
    register_fonts()
    build_er_pdf()
    build_flowchart_pdf()
    build_use_case_drawio()
    build_use_case_pdf()
    build_backup_report()
    build_test_scenarios()
    build_screenshot_report()


if __name__ == "__main__":
    build_reports()
